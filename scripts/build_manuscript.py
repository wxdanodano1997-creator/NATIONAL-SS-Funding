# -*- coding: utf-8 -*-
"""
处理书稿 docx:
  1) 表格序号: 全局连续编号 -> 按章独立编号 (表4.1, 表5.1 ...), 同步更新正文引用
  2) 插入16张词汇联想网络图, 图下方加章节式题注 (图5.1 ...)
"""
import re
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.oxml.ns as ns
from docx.oxml import OxmlElement

SRC = "260620-汉英词汇联想网络增龄研究-参考文献补全版.docx"
OUT = "260621-汉英词汇联想网络增龄研究-图表规范版.docx"
EMBED = "/tmp/embed"
FIGDIR = "网络图-输出"          # 高清原图来源
EMBED_PX = 2200                 # 嵌入分辨率(够清晰且控制docx体积)


def prepare_embeds():
    """把高清网络图按比例缩小后用于嵌入, 控制 docx 体积。"""
    import os
    from PIL import Image
    os.makedirs(EMBED, exist_ok=True)
    for fn in sorted(os.listdir(FIGDIR)):
        if not fn.endswith(".png"):
            continue
        out = os.path.join(EMBED, fn)
        if os.path.exists(out):
            continue
        im = Image.open(os.path.join(FIGDIR, fn)).convert("RGB")
        im.thumbnail((EMBED_PX, EMBED_PX), Image.LANCZOS)
        im.save(out, "PNG", optimize=True)

# ---- 表号: 全局 -> (章, 章内序号) ----
def remap_table(g):
    if 1 <= g <= 6:
        return f"表{4}.{g}"
    if 7 <= g <= 43:
        return f"表5.{g-6}"
    if 44 <= g <= 80:
        return f"表6.{g-43}"
    if 81 <= g <= 122:
        return f"表7.{g-80}"
    return None


def run_spans(p):
    res, pos = [], 0
    for r in p.runs:
        res.append((pos, pos + len(r.text), r))
        pos += len(r.text)
    return res


def apply_span_replace(p, s, e, repl):
    first = True
    for a, b, r in run_spans(p):
        os_, oe = max(s, a), min(e, b)
        if os_ < oe:
            ls, le = os_ - a, oe - a
            if first:
                r.text = r.text[:ls] + repl + r.text[le:]
                first = False
            else:
                r.text = r.text[:ls] + r.text[le:]


def renumber_paragraph(p):
    """把段落内所有 表N 改成 表C.K (跨run安全, 不重复处理已插入文本)。"""
    cursor = 0
    n = 0
    while True:
        text = "".join(r.text for r in p.runs)
        m = re.compile(r"表\s*(\d+)").search(text, cursor)
        if not m:
            break
        g = int(m.group(1))
        repl = remap_table(g)
        if repl is None:
            cursor = m.end()
            continue
        apply_span_replace(p, m.start(), m.end(), repl)
        cursor = m.start() + len(repl)
        n += 1
    return n


def set_caption_font(run, bold=False):
    run.font.size = Pt(12)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(ns.qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(ns.qn("w:eastAsia"), "宋体")
    rf.set(ns.qn("w:ascii"), "Times New Roman")
    rf.set(ns.qn("w:hAnsi"), "Times New Roman")


def main():
    prepare_embeds()
    d = docx.Document(SRC)

    # 1) 表号重排 (正文段落 + 表格单元格段落)
    total = 0
    for p in d.paragraphs:
        total += renumber_paragraph(p)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    total += renumber_paragraph(p)
    print("表号引用更新次数:", total)

    # 2) 插图. 先按当前索引取“下一标题”段落作为插入锚点 (插入用addprevious, 不影响这些对象)
    paras = d.paragraphs
    try:
        style_caption = d.styles["正文但无缩进"]
    except KeyError:
        style_caption = None

    # (锚点段落索引, [(图片文件, 题注bold部分, 题注其余部分)])
    plan = [
        (400, [
            ("初二汉语-网络图.png", "图5.1", "  初二年级汉语词汇联想网络图"),
            ("高二汉语-网络图.png", "图5.2", "  高二年级汉语词汇联想网络图"),
        ]),
        (406, [
            ("初二英语-网络图.png", "图5.3", "  初二年级英语词汇联想网络图"),
            ("高二英语-网络图.png", "图5.4", "  高二年级英语词汇联想网络图"),
        ]),
        (550, [
            ("大一汉语-网络图.png", "图6.1", "  大一年级汉语词汇联想网络图"),
            ("大三汉语-网络图.png", "图6.2", "  大三年级汉语词汇联想网络图"),
        ]),
        (557, [
            ("大一英语-网络图.png", "图6.3", "  大一年级英语词汇联想网络图"),
            ("大三英语-网络图.png", "图6.4", "  大三年级英语词汇联想网络图"),
        ]),
        (689, [
            ("30岁组汉语-网络图.png", "图7.1", "  青年组（30岁组）高校教师汉语词汇联想网络图"),
            ("40岁组汉语-网络图.png", "图7.2", "  中青组（40岁组）高校教师汉语词汇联想网络图"),
            ("50岁组汉语-网络图.png", "图7.3", "  中年组（50岁组）高校教师汉语词汇联想网络图"),
            ("60岁组汉语-网络图.png", "图7.4", "  老年组（60岁组）高校教师汉语词汇联想网络图"),
        ]),
        (695, [
            ("30岁组英语-网络图.png", "图7.5", "  青年组（30岁组）高校教师英语词汇联想网络图"),
            ("40岁组英语-网络图.png", "图7.6", "  中青组（40岁组）高校教师英语词汇联想网络图"),
            ("50岁组英语-网络图.png", "图7.7", "  中年组（50岁组）高校教师英语词汇联想网络图"),
            ("60岁组英语-网络图.png", "图7.8", "  老年组（60岁组）高校教师英语词汇联想网络图"),
        ]),
    ]

    # 校验锚点确实是标题
    for idx, figs in plan:
        print(f"  锚点[{idx}] = {paras[idx].text.strip()[:30]}  插入{len(figs)}张")

    n_fig = 0
    for idx, figs in plan:
        anchor = paras[idx]
        for img, tag, title in figs:
            # 图片段落
            pic_par = d.add_paragraph()
            pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_par.paragraph_format.space_before = Pt(6)
            pic_par.paragraph_format.space_after = Pt(2)
            run = pic_par.add_run()
            run.add_picture(f"{EMBED}/{img}", width=Inches(5.3))
            # 题注段落
            cap_par = d.add_paragraph()
            if style_caption is not None:
                cap_par.style = style_caption
            cap_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_par.paragraph_format.space_after = Pt(10)
            r1 = cap_par.add_run(tag)
            set_caption_font(r1, bold=True)
            r2 = cap_par.add_run(title)
            set_caption_font(r2, bold=False)
            # 移动到锚点之前 (保持顺序)
            anchor._p.addprevious(pic_par._p)
            anchor._p.addprevious(cap_par._p)
            n_fig += 1
    print("插入图片数:", n_fig)

    d.save(OUT)
    print("已保存:", OUT)


if __name__ == "__main__":
    main()
