# -*- coding: utf-8 -*-
"""生成两份补充文档：
   1) 规范参考文献列表（项目成果目录）
   2) 最终成果简介·逐篇介绍版
"""
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, size=12, bold=False, ea="宋体", latin="Times New Roman", color=None):
    run.font.size = Pt(size); run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), ea); rf.set(qn("w:ascii"), latin); rf.set(qn("w:hAnsi"), latin)

def new_doc():
    d = docx.Document()
    for s in d.sections:
        s.left_margin = s.right_margin = Pt(72)
    return d

def para(d, text, size=12, bold=False, align=None, indent=True,
         before=0, after=6, hang=False, color=None, runs=None):
    p = d.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = 1.5
    if align: p.alignment = align
    if hang:
        pf.left_indent = Pt(size * 2); pf.first_line_indent = Pt(-size * 2)
    elif indent:
        pf.first_line_indent = Pt(size * 2)
    if runs:
        for txt, bold_ in runs:
            r = p.add_run(txt); set_font(r, size=size, bold=bold_, color=color)
    else:
        r = p.add_run(text); set_font(r, size=size, bold=bold, color=color)
    return p

def h1(d, t): return para(d, t, size=14, bold=True, indent=False, before=12, after=8)
def h2(d, t): return para(d, t, size=12.5, bold=True, indent=False, before=8, after=4)

# ============ 成果著录（8项，已逐一核对）============
REFS_CN = [
    "张萍. 汉英词汇联想反应分类的共用框架：MMAF的构建及可适性分析[J]. 世界汉语教学, 2026, 40(1): 101–112.（CSSCI）",
    "张萍, 徐雅琛. 语音邻域密度、词性类别、语言水平对二语词汇联想反应的影响[J]. 外语学刊, 2024(1).（CSSCI）",
    "张萍, 马宇晗. 中国英语学习者词汇联想组织模式及网络结构的历时对比研究[J]. 外语教学理论与实践, 2024(3).（CSSCI）",
    "黄旭, 张萍. 我国中老年汉英双语者词汇联想反应类型比较研究[J]. 外国语言文学, 2025(6): 3–13.（北大核心）",
]
REFS_EN = [
    "Wang, X., & Zhang, P. (2024). ‘Did I repeat so many English words?’: Stability of L1 and L2 word association responses over time and across response positions. Lingua, 310, 103804.（SSCI）",
    "Yu, J., Wang, S., Zhang, P., & Chen, T. (2025). The processing of familiar English L2 phrasal verbs in neutral and biased sentence contexts. Frontiers in Psychology, 16, 1528821.（SSCI）",
    "Zhang, P., & Wang, X. (2026). Investigating the bilingual aging lexicon: A network analysis of word associations in Chinese–English bilinguals. Cognitive Science, 50, e70205.（SSCI）",
    "Zhang, P., & Wang, X. (2026). Lexical development in bilingual adolescents: A study of Chinese–English word association networks. Journal of Child Language. Advance online publication. https://doi.org/10.1017/S0305000926100695（SSCI）",
]

# ======================================================================
# 文档一：规范参考文献列表
# ======================================================================
d = new_doc()
para(d, "国家社会科学基金项目成果目录", size=16, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=2)
para(d, "（公开发表论文·参考文献列表）", size=12, bold=False,
     align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=4)
para(d, "说明：本项目共发表学术论文8篇，其中SSCI收录4篇、CSSCI收录3篇、北大核心1篇；"
     "著录格式中文依GB/T 7714，外文依APA。项目正式名称、批准号请依立项通知书填写。",
     size=10.5, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=10,
     color=(0x80, 0x80, 0x80))

h2(d, "一、中文期刊论文")
for i, r in enumerate(REFS_CN, 1):
    para(d, None, size=12, hang=True, after=6,
         runs=[(f"[{i}] ", True), (r, False)])
h2(d, "二、外文期刊论文")
for i, r in enumerate(REFS_EN, 5):
    para(d, None, size=12, hang=True, after=6,
         runs=[(f"[{i}] ", True), (r, False)])

d.save("260709-国社科项目成果-参考文献列表.docx")
print("已保存: 260709-国社科项目成果-参考文献列表.docx")

# ======================================================================
# 文档二：最终成果简介·逐篇介绍版
# ======================================================================
d = new_doc()
para(d, "国家社会科学基金项目最终成果简介（逐篇介绍版）", size=16, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=2)
para(d, "汉英双语心理词库由少至老的发展轨迹：反应类型与网络结构的系统考察", size=13,
     bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=4)
para(d, "（说明：项目正式名称、批准号、负责人、依托单位与起止时间请依立项通知书填入。）",
     size=10.5, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=10,
     color=(0x80, 0x80, 0x80))

h2(d, "（一）研究目的和意义")
para(d, "本项目以汉英双语者心理词库的增龄发展为研究对象，旨在回答一个基础问题：人在不同生命"
     "阶段如何组织与调用汉语母语与英语二语的词汇知识。心理词库并非简单的词表，而是由词形、词义、"
     "句法、语境、搭配与使用经验共同构成的动态网络。随着人口老龄化加深与终身学习体系建设推进，"
     "理解语言能力在青少年、成年与老年阶段的保持、发展与变化，既具有心理语言学与二语习得的理论"
     "价值，也具有服务健康老龄化、老年教育与语言资源建设的现实意义。项目以词汇联想任务为主要"
     "方法，结合反应类型分析与网络分析，从中国汉英双语者的真实语言经验出发，为双语词库的毕生"
     "发展研究提供本土化证据。")
para(d, "具体而言，项目关注四个相互关联的问题：汉语母语词库从青少年到中老年是否保持稳定的语义"
     "主导，其细类反应与网络结构如何随年龄变化；英语二语词库在不同年龄与学习经验阶段如何由形式"
     "依赖走向语义组织，其发展是否与母语同向；汉英两种词库网络在规模、密度、路径、聚集与模块化"
     "等方面呈现何种跨语言差异；词性与具体性等词项属性如何调节不同年龄阶段的联想模式与网络结构。")

h2(d, "（二）成果的主要内容和重要观点")
para(d, "项目共形成学术专著一部与公开发表论文8篇。专著对八个年龄组、两种语言的语义组织与网络"
     "结构作系统整合，8篇论文分别在分类框架、词项属性、历时发展与增龄网络等方面提供实证支撑。"
     "以下按研究逻辑逐篇介绍其主要内容与观点。")
para(d, "总体而言，项目的核心发现可概括为一个判断：汉英双语心理词库的增龄发展不是从不成熟到成熟、"
     "再由成熟到衰退的直线，而是由稳定核心、可变边缘与阶段性重组共同构成的动态过程。母语汉语"
     "较早形成稳固的语义核心，此后主要表现为经验化、语境化与局部联结重心的调整；二语英语具有更长"
     "的发展窗口，青少年阶段以语义化为主，成年早期以搭配化与网络整合为主，成年中后期在长期语言"
     "经验支持下保持较高语义组织水平，同时在形态、母语中介与网络结构上保留更强的可变性。下文"
     "8项成果即从方法框架、词项属性与增龄轨迹等方面为这一判断提供支撑。")

ITEMS = [
    ("成果1　张萍. 汉英词汇联想反应分类的共用框架：MMAF的构建及可适性分析. 世界汉语教学, 2026(1): 101–112.",
     "该文针对词汇联想研究中长期存在的反应分类标准不统一问题，系统回顾一语、二语与双语联想分类"
     "传统，指出既有框架在层级与类别划分、分类定义与可操作性以及汉英分类共用三个方面的不足，进而"
     "构建适用于汉英多年龄段学习/使用者的多层多类分析框架（MMAF），并基于老、中、青三个年龄队列"
     "的汉英联想数据检验其可适性。该框架既保留横组合、纵聚合、非语义等宏观类别，又细分近义、反义、"
     "上下义、同级、修饰、动作、地点、工具、语境、音形、母语中介、派生、屈折等类别，是本项目的"
     "方法论基础，保证了不同年龄、语言与词项属性之间的比较具有共同尺度。"),
    ("成果2　张萍, 徐雅琛. 语音邻域密度、词性类别、语言水平对二语词汇联想反应的影响. 外语学刊, 2024(1).",
     "该文采用“听—写”词汇联想任务，考察语音邻域密度、词性与二语水平对中国英语学习者联想反应的"
     "影响。研究发现：学习者均以语义反应为主，语音反应比例与语言水平呈负相关；不同词性主要影响"
     "语义联想，对语音反应影响不显著；语音邻域密度仅在高水平者的语义反应及低水平者的名词反应上"
     "表现出作用，部分支持邻域激活模型。该成果说明二语词库在发展中逐步减弱对语音形式线索的依赖、"
     "转向更稳定的语义组织，并揭示词性与语音邻域等词项属性对联想路径的调节作用。"),
    ("成果3　张萍, 马宇晗. 中国英语学习者词汇联想组织模式及网络结构的历时对比研究. 外语教学理论与实践, 2024(3).",
     "该文将词汇联想与网络科学结合，历时比较中国英语学习者在不同学习阶段的二语词库组织。研究"
     "表明，随学习年限与语言经验增加，学习者的二语联想不仅在反应类型上趋于语义化，网络结构上"
     "也出现密度、聚集性与通达效率的变化，高年级学习者的网络连通性更高、路径更短。该成果把外语"
     "学习刻画为不断增加节点、建立联结并优化通达路径的网络建构过程，为本项目采用网络指标分析"
     "增龄轨迹提供了直接的方法参照。"),
    ("成果4　Yu, J., Wang, S., Zhang, P., & Chen, T. (2025). The processing of familiar English L2 "
     "phrasal verbs in neutral and biased sentence contexts. Frontiers in Psychology, 16, 1528821.",
     "该文结合眼动追踪与视觉词汇搜索任务，考察二语学习者对熟悉英语短语动词在中性、字面义偏向与"
     "引申义偏向语境中的加工。研究发现，熟悉短语动词的意义激活并不遵循固定的字面义优先或引申义"
     "优先路径，而表现出明显的时间动态性与语境依赖性；在熟悉度得到控制的条件下，中、高级学习者"
     "之间未见显著差异。该成果说明心理词库中并非只有孤立词项，短语动词、搭配与语块等多词单位同样"
     "具有整体化表征，并受语境与使用经验调节。"),
    ("成果5　Wang, X., & Zhang, P. (2024). ‘Did I repeat so many English words?’: Stability of L1 and "
     "L2 word association responses over time and across response positions. Lingua, 310, 103804.",
     "该文以连续联想任务与重复测量设计，考察一语与二语联想反应的稳定性。研究发现，二语反应在两次"
     "测试间的重复率高于一语，反应位置显著影响稳定性，首位反应比后续反应更稳定，且个体的一语与"
     "二语稳定性之间存在相关。该成果从短时历时角度说明联想反应兼具稳定性与变异性；较高的二语"
     "重复率并不简单意味着二语更成熟，而可能反映其可用语义路径较为集中，为理解词库中核心联结与"
     "边缘联结的关系提供了参照。"),
    ("成果6　Zhang, P., & Wang, X. (2026). Lexical development in bilingual adolescents: A study of "
     "Chinese–English word association networks. Journal of Child Language（在线发表）.",
     "该文考察初二与高二学生的汉英词汇联想网络。研究发现，从初二到高二，汉语网络在保持整体结构"
     "与小世界特征的同时组内趋同增强；英语网络则出现规模扩张、平均最短路径延长、局部聚类与模块化"
     "提高，反映二语词汇系统在青少年阶段的快速增长与重组；两种语言的网络差距随年级增长而缩小。"
     "该成果为本项目青少年阶段提供了关键证据，说明青春期并非词汇发展的终点，而是双语词库结构"
     "继续重组的重要阶段。"),
    ("成果7　Zhang, P., & Wang, X. (2026). Investigating the bilingual aging lexicon: A network "
     "analysis of word associations in Chinese–English bilinguals. Cognitive Science, 50, e70205.",
     "该文以网络分析考察四个成人年龄组汉英双语者的词汇联想网络。研究发现，汉语母语网络从成年早期"
     "到中年全局连通性增强、局部聚集下降、通达效率提高，中年之后趋于稳定，老年阶段出现效率下降与"
     "模块化增强；英语二语网络在成年早中期方向与汉语相似，但结构偏移幅度更大，且在整个成年期的"
     "连通性与模块化程度均低于汉语。该成果提出双语老化词库的“语言非对称”图景，说明强势母语具有"
     "较高的结构稳定性，相对弱势的二语更易受年龄、经验与加工资源变化影响。"),
    ("成果8　黄旭, 张萍. 我国中老年汉英双语者词汇联想反应类型比较研究. 外国语言文学, 2025(6): 3–13.",
     "该文聚焦中年与老年高校英语教师，比较其汉英词汇联想反应类型。研究发现，两组双语者的汉语"
     "联想均以语义关系为主并倾向横组合，但在“音形”“其他”“上下义”“同级”等细类上存在年龄"
     "差异；英语联想同样以语义为主但更偏纵聚合，尤在“派生”“屈折”“母语中介”和“补充”反应上"
     "表现不同；年龄与语言类型仅在横组合的“工具”反应上呈现交互效应。该成果为本项目成年中后期"
     "样本提供了直接证据，说明中老年双语词库并未发生整体性语义退化，而是在不同语言与不同细类"
     "联结上呈现局部调整。"),
]
for cite, body in ITEMS:
    para(d, cite, size=12, bold=True, indent=False, before=6, after=2)
    para(d, body, size=12, indent=True, after=6)

para(d, "综合来看，8项成果并非彼此分散的个案，而是围绕“汉英双语心理词库如何随年龄与经验发展”"
     "这一核心问题逐层展开、相互支撑：MMAF框架解决分类尺度问题，语音邻域与词性研究、短语动词"
     "加工研究说明词项属性与多词单位如何影响联想路径，历时对比与青少年网络研究刻画二语词库随"
     "学习阶段的发展，重复测量研究揭示联想反应的稳定性与变异性，成人网络研究与中老年反应类型"
     "研究把考察推进到成年后期与老年阶段。由此，成果在研究对象上覆盖由少至老，在方法上兼顾"
     "反应分类与网络分析，在理论解释上形成从局部机制到整体轨迹的互补关系。", before=6)

para(d, "在此基础上，项目形成三点重要判断。第一，母语与二语共享概念基础，但发展轨迹并不相同："
     "母语汉语更早形成稳定语义网络，二语英语更受学习经验、使用频率与任务环境影响。第二，年龄"
     "增长并不等于词库整体退化，尤其在高教育、高语言经验群体中，成年后期词库仍能保持语义组织，"
     "并通过语境化、模块化与策略性联想表现出新的组织特点。第三，词汇能力的核心不只是词汇量，"
     "而是词与词之间能否形成可通达、可迁移、可用于理解与表达的关系。这些判断对外语词汇教学、"
     "终身学习与老年语言教育具有启示：词汇学习宜从孤立记忆转向关系建构，语言老化研究宜从缺损"
     "叙事转向“保持与重组并存”的解释。")

para(d, "作为项目的集成性成果，学术专著《心理词库联想网络增龄化发展轨迹研究》将上述发现纳入统一"
     "框架，系统呈现由少至老三个阶段、两种语言在语义组织与网络结构上的发展轨迹及其影响因素——"
     "包括汉语语义框架较早稳固而细类随年龄再分配、英语由形式依赖走向语义化与搭配化、汉英之间"
     "共享与分化并存、词性与具体性持续调节联想路径等，是对上述系列论文的整体提炼与理论深化。", before=6)

h2(d, "（三）成果的主要价值和影响")
para(d, "项目价值集中于三方面。理论上，将双语心理词库研究与毕生发展、认知老化与网络科学结合，"
     "深化了对母语与二语非对称发展、词库可塑性及语言经验调节作用的认识。方法上，提出并应用汉英"
     "词汇联想反应分类框架，形成语义分析与网络分析结合的研究路径，为后续研究提供了可复用的分析"
     "工具与数据基础。实践上，研究结论可为青少年英语词汇教学、大学阶段二语词汇深度发展、中老年"
     "语言学习、认知健康研究与汉英语义资源建设提供参考。")
para(d, "从影响看，成果已在国内外多层次期刊发表，覆盖中文核心外语类期刊与国际心理语言学、认知"
     "科学、儿童语言发展期刊，表明该议题能够进入多个学术讨论领域；项目同时形成了可继续扩展的"
     "研究平台——汉英词汇联想数据可作为常模建设与跨研究比较的基础，反应分类框架与网络分析流程"
     "可迁移至其他语言组合、其他年龄群体以及教育或临床场景。就此而言，项目成果不止于完成一次"
     "结项，而是为后续围绕汉英双语心理词库、终身语言学习与健康老龄化的连续研究奠定了基础。")
para(d, "从现实层面看，项目对语言教育与老年学习具有较直接的启示。对于青少年与大学生，词汇教学"
     "宜从词义解释、背诵与测试，逐步转向语义场、搭配、语块、语境与跨语言关系的建构；对于中老年"
     "学习者，语言学习不应被视为单纯补偿衰退的活动，而应被理解为调动已有经验、维持交流与促进"
     "社会参与的认知活动。项目关于健康人群词库结构的发现，也可为今后开展正常老化与异常退化的"
     "比较研究，以及认知筛查与老年语言教育项目设计提供参照。")

d.save("260709-最终成果简介-逐篇介绍版.docx")
print("已保存: 260709-最终成果简介-逐篇介绍版.docx")

import re
def wc(doc):
    s="".join(p.text for p in doc.paragraphs)
    return len(re.findall(r'[一-鿿]', s))+len(re.findall(r'[A-Za-z]+', s))+len(re.findall(r'\d+', s))
print("逐篇介绍版 近似字数:", wc(d))
